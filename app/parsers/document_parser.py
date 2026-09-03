from __future__ import annotations

import logging
from pathlib import Path

import pymupdf
from docx import Document
from openpyxl import load_workbook

from app.models import PageText


LOGGER = logging.getLogger(__name__)


def parse_document(path: Path) -> list[PageText]:
    """Extract page/sheet-aware text without executing macros or embedded content."""
    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            return _parse_pdf(path)
        if suffix == ".docx":
            return _parse_docx(path)
        if suffix == ".xlsx":
            return _parse_xlsx(path)
        if suffix == ".txt":
            return [PageText(1, path.read_text(encoding="utf-8", errors="replace"))]
        raise ValueError(f"不支持解析 {suffix}")
    except Exception:
        LOGGER.exception("文件解析失败：%s", path.name)
        raise


def _parse_pdf(path: Path) -> list[PageText]:
    pages: list[PageText] = []
    with pymupdf.open(path) as document:
        if document.needs_pass:
            raise ValueError("PDF 已加密，无法解析")
        for number, page in enumerate(document, start=1):
            text = page.get_text("text", sort=True).strip()
            table_text = _extract_pdf_tables(page)
            if table_text:
                text = f"{text}\n\n{table_text}".strip()
            pages.append(PageText(number, text))
    return pages


def _extract_pdf_tables(page: pymupdf.Page) -> str:
    """Keep row relationships that plain PDF text extraction usually destroys."""
    try:
        tables = page.find_tables().tables
    except Exception as exc:
        LOGGER.warning("PDF 表格识别失败，保留普通文本：%s", type(exc).__name__)
        return ""
    lines: list[str] = []
    for table in tables:
        for row in table.extract():
            cells = [" ".join(str(cell).split()) if cell is not None else "" for cell in row]
            if any(cells):
                lines.append("[TABLE_ROW] " + " || ".join(cells))
    return "\n".join(lines)


def _parse_docx(path: Path) -> list[PageText]:
    document = Document(path)
    lines = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        lines.extend("\t".join(cell.text for cell in row.cells) for row in table.rows)
    return [PageText(1, "\n".join(item for item in lines if item.strip()))]


def _parse_xlsx(path: Path) -> list[PageText]:
    workbook = load_workbook(path, read_only=True, data_only=True, keep_links=False)
    pages: list[PageText] = []
    for number, sheet in enumerate(workbook.worksheets, start=1):
        rows = ["\t".join("" if value is None else str(value) for value in row) for row in sheet.iter_rows(values_only=True)]
        pages.append(PageText(number, f"工作表：{sheet.title}\n" + "\n".join(rows)))
    workbook.close()
    return pages
