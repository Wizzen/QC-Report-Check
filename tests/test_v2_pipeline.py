from __future__ import annotations

from io import BytesIO
import json
from pathlib import Path
from threading import Barrier, Lock
import time
from unittest.mock import patch

import pymupdf
import pytest
from docx import Document

from app.auditing.v2_service import ReviewService, _expert_document_chunks
from app.auditing.expert_review import ExpertDocument, build_expert_prompt
from app.database import ReviewDatabase
from app.integrations import ConfigStore
from app.models import Finding, PageText


class NamedBytesIO(BytesIO):
    def __init__(self, data: bytes, name: str):
        super().__init__(data)
        self.name = name


def _pdf(text: str, name: str) -> NamedBytesIO:
    document = pymupdf.open()
    page = document.new_page()
    page.insert_text((72, 72), text, fontname="china-s")
    payload = document.tobytes()
    document.close()
    return NamedBytesIO(payload, name)


def _docx(text: str, name: str) -> NamedBytesIO:
    document = Document()
    document.add_paragraph(text)
    payload = BytesIO()
    document.save(payload)
    return NamedBytesIO(payload.getvalue(), name)


def _service(tmp_path: Path) -> ReviewService:
    db = ReviewDatabase(tmp_path / "review.db")
    store = ConfigStore(db, tmp_path / "service.key")
    return ReviewService(db, store, tmp_path / "uploads", tmp_path / "standards", tmp_path / "vectors")


def test_review_requires_supplier_file_before_creating_batch(tmp_path: Path) -> None:
    service = _service(tmp_path)

    with pytest.raises(ValueError, match="至少需要"):
        service.create_review(None, [], [], [])

    assert service.db.one("SELECT COUNT(*) count FROM review_batches")["count"] == 0


def test_each_upload_creates_an_independent_event_and_document_copy(tmp_path: Path) -> None:
    service = _service(tmp_path)
    template_id = service.db.one("SELECT id FROM audit_templates WHERE is_default=1")["id"]

    first_batch = service.create_review(
        template_id, [], [_pdf("Report No R-1", "same-name.pdf")], [], llm_concurrency=2
    )
    second_batch = service.create_review(template_id, [], [_pdf("Report No R-2", "same-name.pdf")], [])

    assert first_batch != second_batch
    first_document = service.db.one(
        "SELECT d.id,d.stored_path FROM batch_documents bd JOIN documents d ON d.id=bd.document_id WHERE bd.batch_id=?",
        (first_batch,),
    )
    second_document = service.db.one(
        "SELECT d.id,d.stored_path FROM batch_documents bd JOIN documents d ON d.id=bd.document_id WHERE bd.batch_id=?",
        (second_batch,),
    )
    assert first_document["id"] != second_document["id"]
    assert first_document["stored_path"] != second_document["stored_path"]
    assert first_batch in first_document["stored_path"]
    assert second_batch in second_document["stored_path"]
    assert service.db.one("SELECT template_snapshot FROM review_batches WHERE id=?", (first_batch,))["template_snapshot"]
    assert service.db.one("SELECT template_snapshot FROM review_batches WHERE id=?", (second_batch,))["template_snapshot"]
    assert service.db.one("SELECT review_mode FROM review_batches WHERE id=?", (first_batch,))["review_mode"] == "adaptive"
    assert service.db.one("SELECT llm_concurrency FROM review_batches WHERE id=?", (first_batch,))["llm_concurrency"] == 2


def test_full_local_rule_path_uses_basis_not_supplier_archive(tmp_path: Path) -> None:
    service = _service(tmp_path)
    basis = NamedBytesIO("抗拉强度 >= 500 MPa".encode("utf-8"), "采购技术要求.txt")
    with patch("app.integrations.clients.EmbeddingClient.embed") as embed:
        basis_id = service.import_basis(basis, "technical")
        supplier = _docx(
            "Material Test Certificate  Supplier Quality Record  "
            "Heat No: H20260824  Material Grade: Q355B  Rm: 480 MPa  "
            "This certificate contains verified inspection results.",
            "supplier.docx",
        )
        batch_id = service.create_review(None, [basis_id], [supplier], [])
        service.process_batch(batch_id)
    embed.assert_not_called()

    batch = service.db.one("SELECT * FROM review_batches WHERE id=?", (batch_id,))
    findings = service.db.query("SELECT * FROM findings WHERE batch_id=?", (batch_id,))
    assert batch and batch["status"] == "completed"
    assert any(row["item"] == "抗拉强度" and row["severity"] == "Major" for row in findings)
    selected_basis = service.db.query(
        "SELECT d.library_code FROM batch_documents bd JOIN documents d ON d.id=bd.document_id "
        "WHERE bd.batch_id=? AND bd.role IN ('selected_basis','supplemental_basis')",
        (batch_id,),
    )
    assert selected_basis and {row["library_code"] for row in selected_basis} == {"basis"}
    assert service.db.one("SELECT index_status FROM documents WHERE id=?", (basis_id,))["index_status"] == "not_used"


def test_llm_explanations_are_batched_instead_of_one_call_per_finding(tmp_path: Path) -> None:
    service = _service(tmp_path)
    service.config_store.save({"llm_model": "test-model"})
    batch_id = service.db.create_batch(None)
    findings = [Finding("数值不符合", "Major", f"item-{index}", "old", actual="1", requirement=">=2")
                for index in range(3)]
    response = {"items": [
        {"index": index, "reason": f"reason-{index}", "suggestion": f"fix-{index}", "confidence": 0.9}
        for index in range(3)
    ]}

    with patch("app.auditing.v2_service.LLMClient.generate_json", return_value=response) as generate:
        result = service._explain(batch_id, findings)

    assert generate.call_count == 1
    assert [item.description for item in result] == ["reason-0", "reason-1", "reason-2"]
    batch = service.db.one("SELECT activity,resource,heartbeat_at FROM review_batches WHERE id=?", (batch_id,))
    assert "批量解释" in batch["activity"]
    assert "LLM" in batch["resource"]
    assert batch["heartbeat_at"]


def test_fastener_template_alias_keeps_proven_report_checks(tmp_path: Path) -> None:
    """Regression: the user's existing template name must not disable fastener checks."""
    service = _service(tmp_path)
    template_id = service.db.execute(
        """INSERT INTO audit_templates(name,description,required_document_types,required_items,review_instructions,
           enabled,is_default,created_at) VALUES(?,?,?,?,?,1,0,datetime('now'))""",
        ("维修备件紧固件审核", "", "[]", "[]", "核验 WDC、COC/MTR 以及 Sample/Pass 表格"),
    )
    batch_id = service.db.create_batch(template_id)
    document_id = service.db.add_document(
        library="supplier", kind="supplier", original_name="Q0045 5305-859240(MTR).pdf",
        stored_path=str(tmp_path / "same-report.pdf"), sha256="same-report",
    )
    page_text = """QUALITY CERTIFICATE
CHEMICAL COMPOSITION(%)
Mechanical Properties
Dimensions Of SPEC
[TABLE_ROW] Head Height || 10.18-9.82 || 9.91-10.00 || 20 || 19
[TABLE_ROW] Thread Length || min 38. || 40.00-40.00 || 20 || 17
[TABLE_ROW] HV(2)>=HV(1)-30 || G 0.015max || 6 || 6
"""
    service.db.execute(
        "UPDATE documents SET page_text=?,parse_status='completed' WHERE id=?",
        (json.dumps([{"page": 1, "text": page_text}], ensure_ascii=False), document_id),
    )
    service.db.attach_document(batch_id, document_id, "supplier", 4)

    findings = service._audit(batch_id)

    assert any(item.item == "Head Height" and item.severity == "Major" for item in findings)
    assert any(item.item == "Thread Length" and item.severity == "Major" for item in findings)


def test_llm_failure_never_claims_zero_problems(tmp_path: Path) -> None:
    service = _service(tmp_path)
    service.config_store.save({"llm_model": "test-model"})
    template_id = service.db.execute(
        """INSERT INTO audit_templates(name,description,required_document_types,required_items,review_instructions,
           enabled,is_default,created_at) VALUES(?,?,?,?,?,1,0,datetime('now'))""",
        ("失败隔离模板", "", "[]", json.dumps(["检查报告编号"], ensure_ascii=False), "只依据原文"),
    )
    batch_id = service.db.create_batch(template_id)
    document_id = service.db.add_document(
        library="supplier", kind="supplier", original_name="unknown.pdf",
        stored_path=str(tmp_path / "unknown.pdf"), sha256="unknown",
    )
    service.db.execute(
        "UPDATE documents SET page_text=?,parse_status='completed' WHERE id=?",
        (json.dumps([{"page": 1, "text": "Readable supplier inspection report with no deterministic mismatch."}]), document_id),
    )
    service.db.attach_document(batch_id, document_id, "supplier", 4)

    with patch("app.auditing.v2_service.LLMClient.generate_json", side_effect=RuntimeError("offline")):
        findings = service._expert_review(batch_id, [])

    assert len(findings) == 1
    assert findings[0].severity == "Review"
    assert findings[0].item == "部分独立规则未完成"


def test_each_template_item_calls_llm_once_and_is_aggregated(tmp_path: Path) -> None:
    service = _service(tmp_path)
    service.config_store.save({"llm_model": "test-model"})
    tasks = ["检查炉号", "检查客户签字", "检查报告编号"]
    template_id = service.db.execute(
        """INSERT INTO audit_templates(name,description,required_document_types,required_items,review_instructions,
           enabled,is_default,created_at) VALUES(?,?,?,?,?,1,0,datetime('now'))""",
        ("逐条模板", "", "[]", json.dumps(tasks, ensure_ascii=False), "只依据原文"),
    )
    batch_id = service.db.create_batch(template_id, "deep")
    document_id = service.db.add_document(
        library="supplier", kind="supplier", original_name="MTR.pdf",
        stored_path=str(tmp_path / "MTR.pdf"), sha256="independent-rules",
    )
    service.db.execute(
        "UPDATE documents SET page_text=?,parse_status='completed' WHERE id=?",
        (json.dumps([{"page": 1, "text": "Heat No: H123\nReport No: R456"}], ensure_ascii=False), document_id),
    )
    service.db.attach_document(batch_id, document_id, "supplier", 4)
    responses = [
        {"result": "合格", "conclusion": "炉号存在", "source_file": "MTR.pdf", "page": 1,
         "evidence": "Heat No: H123", "confidence": 0.95},
        {"result": "不适用", "conclusion": "本报告无需客户签字", "confidence": 0.9},
        {"result": "存疑", "conclusion": "报告编号需复核", "source_file": "MTR.pdf", "page": 1,
         "evidence": "Report No: R456", "logic": "格式不明确", "confidence": 0.72},
    ]

    with patch("app.auditing.v2_service.LLMClient.generate_json", side_effect=responses) as generate:
        findings = service._expert_review(batch_id, [])

    assert generate.call_count == len(tasks)
    evaluations = service.db.query(
        "SELECT task_name,status FROM rule_evaluations WHERE batch_id=? ORDER BY task_index", (batch_id,)
    )
    assert [(row["task_name"], row["status"]) for row in evaluations] == [
        (tasks[0], "合格"), (tasks[1], "不适用"), (tasks[2], "存疑")]
    assert [finding.item for finding in findings] == ["检查报告编号"]


def test_template_rule_tasks_are_submitted_concurrently(tmp_path: Path) -> None:
    service = _service(tmp_path)
    service.config_store.save({"llm_model": "test-model", "llm_concurrency": 1})
    tasks = ["任务一", "任务二", "任务三"]
    template_id = service.db.execute(
        """INSERT INTO audit_templates(name,description,required_document_types,required_items,review_instructions,
           enabled,is_default,created_at) VALUES(?,?,?,?,?,1,0,datetime('now'))""",
        ("并发模板", "", "[]", json.dumps(tasks, ensure_ascii=False), "只依据原文"),
    )
    batch_id = service.db.create_batch(template_id, llm_concurrency=3)
    document_id = service.db.add_document(
        library="supplier", kind="supplier", original_name="MTR.pdf",
        stored_path=str(tmp_path / "MTR.pdf"), sha256="parallel-rules",
    )
    service.db.execute(
        "UPDATE documents SET page_text=?,parse_status='completed' WHERE id=?",
        (json.dumps([{"page": 1, "text": "Inspection Report"}], ensure_ascii=False), document_id),
    )
    service.db.attach_document(batch_id, document_id, "supplier", 4)
    state_lock = Lock()
    active = 0
    max_active = 0

    def concurrent_response(*_args, **_kwargs):
        nonlocal active, max_active
        with state_lock:
            active += 1
            max_active = max(max_active, active)
        time.sleep(0.03)
        with state_lock:
            active -= 1
        return {"result": "不适用", "conclusion": "测试完成", "confidence": 0.9}

    with patch("app.auditing.v2_service.LLMClient.generate_json", side_effect=concurrent_response) as generate:
        service._expert_review(batch_id, [])

    assert generate.call_count == 3
    assert max_active == 2
    assert all(call.kwargs["timeout_seconds"] == 90 for call in generate.call_args_list)
    assert all(call.kwargs["max_tokens"] == 220 for call in generate.call_args_list)
    assert all(call.kwargs["retries"] == 1 for call in generate.call_args_list)
    assert all(call.kwargs["thinking"] is False for call in generate.call_args_list)
    assert [row["status"] for row in service.db.query(
        "SELECT status FROM rule_evaluations WHERE batch_id=? ORDER BY task_index", (batch_id,))] == [
            "不适用", "不适用", "不适用"]


def test_adaptive_mode_calls_each_rule_once_and_downgrades_unlocated_evidence(tmp_path: Path) -> None:
    service = _service(tmp_path)
    service.config_store.save({"llm_model": "qwen3.5-test", "llm_concurrency": 2})
    template_id = service.db.execute(
        """INSERT INTO audit_templates(name,description,required_document_types,required_items,review_instructions,
           enabled,is_default,created_at) VALUES(?,?,?,?,?,1,0,datetime('now'))""",
        ("自适应模板", "", "[]", json.dumps(["检查报告编号"], ensure_ascii=False), "只依据原文"),
    )
    batch_id = service.db.create_batch(template_id, "adaptive")
    document_id = service.db.add_document(
        library="supplier", kind="supplier", original_name="MTR.pdf",
        stored_path=str(tmp_path / "MTR.pdf"), sha256="adaptive-mode",
    )
    service.db.execute(
        "UPDATE documents SET page_text=?,parse_status='completed' WHERE id=?",
        (json.dumps([{"page": 1, "text": "Report No: R456"}], ensure_ascii=False), document_id),
    )
    service.db.attach_document(batch_id, document_id, "supplier", 4)
    response = {"r": "不合格", "c": "需要复核", "f": "MTR.pdf", "p": 1,
                "e": "模型转述而非原页逐字证据", "q": 0.8}

    with patch("app.auditing.v2_service.LLMClient.generate_json", return_value=response) as generate:
        findings = service._expert_review(batch_id, [])

    assert len(findings) == 1
    assert findings[0].severity == "Review"
    assert findings[0].source_text == ""
    assert findings[0].metadata["evidence_type"] == "unlocated"
    assert generate.call_count == 1
    assert generate.call_args.kwargs == {
        "retries": 1, "timeout_seconds": 90, "max_tokens": 220, "thinking": False,
    }
    assert "每条规则只调用一次模型" in generate.call_args.args[0]
    assert service.db.one(
        "SELECT status FROM rule_evaluations WHERE batch_id=?", (batch_id,)
    )["status"] == "存疑"


def test_batch_uses_template_snapshot_after_template_is_edited(tmp_path: Path) -> None:
    service = _service(tmp_path)
    service.config_store.save({"llm_model": "test-model", "llm_concurrency": 2})
    template_id = service.db.execute(
        """INSERT INTO audit_templates(name,description,required_document_types,required_items,review_instructions,
           enabled,is_default,created_at) VALUES(?,?,?,?,?,1,0,datetime('now'))""",
        ("快照模板", "", "[]", json.dumps(["上传时任务A", "上传时任务B"], ensure_ascii=False), "上传时说明"),
    )
    batch_id = service.db.create_batch(template_id)
    service.db.execute(
        "UPDATE audit_templates SET required_items=?,review_instructions=? WHERE id=?",
        (json.dumps(["后来任务"], ensure_ascii=False), "后来说明", template_id),
    )
    document_id = service.db.add_document(
        library="supplier", kind="supplier", original_name="MTR.pdf",
        stored_path=str(tmp_path / "MTR.pdf"), sha256="snapshot-rules",
    )
    service.db.execute(
        "UPDATE documents SET page_text=?,parse_status='completed' WHERE id=?",
        (json.dumps([{"page": 1, "text": "Inspection Report"}], ensure_ascii=False), document_id),
    )
    service.db.attach_document(batch_id, document_id, "supplier", 4)

    with patch("app.auditing.v2_service.LLMClient.generate_json",
               return_value={"result": "不适用", "conclusion": "完成"}) as generate:
        service._expert_review(batch_id, [])

    assert generate.call_count == 2
    assert [row["task_name"] for row in service.db.query(
        "SELECT task_name FROM rule_evaluations WHERE batch_id=? ORDER BY task_index", (batch_id,))] == [
            "上传时任务A", "上传时任务B"]


def test_one_failed_rule_does_not_stop_later_rules(tmp_path: Path) -> None:
    service = _service(tmp_path)
    service.config_store.save({"llm_model": "test-model"})
    tasks = ["任务一", "任务二", "任务三"]
    template_id = service.db.execute(
        """INSERT INTO audit_templates(name,description,required_document_types,required_items,review_instructions,
           enabled,is_default,created_at) VALUES(?,?,?,?,?,1,0,datetime('now'))""",
        ("失败隔离", "", "[]", json.dumps(tasks, ensure_ascii=False), "只依据原文"),
    )
    batch_id = service.db.create_batch(template_id)
    document_id = service.db.add_document(
        library="supplier", kind="supplier", original_name="MTR.pdf",
        stored_path=str(tmp_path / "MTR.pdf"), sha256="failure-isolation",
    )
    service.db.execute(
        "UPDATE documents SET page_text=?,parse_status='completed' WHERE id=?",
        (json.dumps([{"page": 1, "text": "Inspection Report"}], ensure_ascii=False), document_id),
    )
    service.db.attach_document(batch_id, document_id, "supplier", 4)
    responses = [
        {"result": "合格", "conclusion": "完成"}, RuntimeError("temporary offline"),
        {"result": "不适用", "conclusion": "完成"},
    ]

    with patch("app.auditing.v2_service.LLMClient.generate_json", side_effect=responses) as generate:
        findings = service._expert_review(batch_id, [])

    assert generate.call_count == 3
    assert [row["status"] for row in service.db.query(
        "SELECT status FROM rule_evaluations WHERE batch_id=? ORDER BY task_index", (batch_id,))] == [
            "合格", "调用失败", "不适用"]
    assert any(finding.item == "部分独立规则未完成" for finding in findings)


def test_long_report_is_split_for_small_context_models(tmp_path: Path) -> None:
    document = ExpertDocument(
        "long-report.pdf", tmp_path / "long-report.pdf",
        [PageText(page, f"page-{page} " + ("inspection evidence " * 700)) for page in range(1, 5)],
    )

    chunks = _expert_document_chunks([document])

    assert len(chunks) >= 4
    assert {page.page for chunk in chunks for item in chunk for page in item.pages} == {1, 2, 3, 4}
    assert all(len(build_expert_prompt(chunk, "", [])) < 24_000 for chunk in chunks)
